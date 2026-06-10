# ============================================================
# NOME DO SCRIPT: relatorio_semanal_besst.py
# OBJETIVO: Gerar relatório semanal em Word (.docx) com ranking
#           TOP 10 BESST e análise aprofundada das TOP 5 ações,
#           combinando dados quantitativos e análise via API Gemini.
# BIBLIOTECAS: python-docx, yfinance, requests, pandas, google-genai
# EXECUTAR:    python relatorio_semanal_besst.py
# INSTALAR:    pip install python-docx yfinance requests pandas google-genai
# ============================================================

# --- SEÇÃO 1: Importações ---
import os, time, json, warnings, math
from datetime import datetime, timedelta
from io import StringIO

import pandas as pd
import requests
import yfinance as yf

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from google import genai

warnings.filterwarnings("ignore")


# --- SEÇÃO 2: Configurações ---

PASTA_RESULTADOS  = "resultados"
PASTA_RELATORIOS  = "resultados/relatorios"
PASTA_CACHE       = "dados/fundamentos"
MODELO_GEMINI     = "gemini-2.5-flash"

FILTRO_LIQUIDEZ_MINIMA = 500_000
FILTRO_PL_MINIMO       = 0
FILTRO_PL_MAXIMO       = 40
FILTRO_DY_MINIMO       = 0.06
FILTRO_PVPA_MAXIMO     = 5.0
FILTRO_BESST_ATIVO     = True

TICKERS_BESST = {
    "Bancos":     ["BBAS3","BBDC3","BBDC4","ITUB4","ITSA4","SANB11","BPAC11"],
    "Energia":    ["ELET3","ELET6","CMIG4","CPFE3","CPLE6","EGIE3","ENEV3",
                   "ENGI11","EQTL3","TAEE11","ISAE4"],
    "Saneamento": ["SBSP3"],
    "Seguros":    ["BBSE3","IRBR3","PSSA3"],
    "Telecom":    ["VIVT3","TIMS3"],
}
TODOS_TICKERS_BESST = [t for lista in TICKERS_BESST.values() for t in lista]

PESOS = {
    "dy": 0.30, "roe": 0.25, "pl": 0.20, "pvpa": 0.15, "div_ebitda": 0.10,
}

COR_TITULO    = RGBColor(0x1F, 0x49, 0x7D)
COR_SUBTITULO = RGBColor(0x2E, 0x75, 0xB6)
COR_DESTAQUE  = RGBColor(0x70, 0xAD, 0x47)
COR_RISCO     = RGBColor(0xFF, 0x00, 0x00)

TOP_RANKING    = 10
TOP_ANALISE    = 5
ANOS_HISTORICO = 5

TEXTO_INTRODUCAO = """Este relatório foi gerado automaticamente por meio de coleta, processamento e consolidação de dados públicos disponíveis em portais especializados do mercado financeiro brasileiro, como Fundamentus, StatusInvest, Yahoo Finance e a base de dados oficial da CVM (Comissão de Valores Mobiliários). Nenhuma informação utilizada nesta análise é proprietária ou restrita — todas as fontes consultadas são de acesso público e gratuito."""

TEXTO_OBJETIVO = """O propósito deste relatório é reunir, em um único documento estruturado, os principais indicadores fundamentalistas de empresas listadas na B3 (Bolsa de Valores do Brasil), facilitando a leitura e a comparação entre ativos para o investidor pessoa física. A consolidação dessas informações tem caráter exclusivamente educacional e informativo, buscando democratizar o acesso a dados que, embora públicos, estão dispersos em diferentes fontes e formatos.

A filosofia que norteia a seleção dos ativos analisados é a de acumulação de longo prazo: adquirir participações em empresas sólidas, com histórico consistente de geração de caixa e distribuição de proventos, e mantê-las pelo maior prazo possível. Essa abordagem — conhecida no mercado como buy and hold — difere fundamentalmente das estratégias especulativas de curto prazo, que buscam lucro na oscilação de preços por meio de compras e vendas frequentes de papéis. Enquanto o especulador depende do comportamento do mercado no curtíssimo prazo, o investidor de longo prazo constrói renda passiva real e crescente ao longo dos anos, tornando-se progressivamente menos dependente do trabalho ativo para sustentar seu padrão de vida."""

TEXTO_METODOLOGIA = """Os dados foram coletados de forma automatizada via scripts Python, utilizando bibliotecas de acesso a APIs públicas e técnicas de raspagem de dados (web scraping) em portais de informação financeira. Os indicadores apresentados — como P/L, P/VP, Dividend Yield, ROE, Dívida Líquida/EBITDA e histórico de dividendos — foram calculados ou extraídos diretamente das fontes originais, sem alteração de conteúdo. Cada seção do relatório indica a fonte dos dados e a data de coleta, permitindo rastreabilidade e verificação independente pelo leitor."""

TEXTO_LIMITACOES = """Por se tratar de coleta automatizada de fontes externas, os dados podem apresentar defasagem em relação à publicação oficial mais recente. Indicadores baseados em resultados trimestrais e anuais refletem as demonstrações financeiras já divulgadas pelas empresas à CVM até a data de geração deste relatório. Recomenda-se sempre verificar os dados diretamente nas fontes originais antes de qualquer utilização."""

TEXTO_AVISO = """As informações contidas neste relatório têm finalidade exclusivamente informativa e educacional. Nenhum conteúdo aqui apresentado representa, de qualquer forma, recomendação de compra, venda ou manutenção de valores mobiliários. A análise de indicadores fundamentalistas é apenas um dos múltiplos fatores que devem ser considerados em uma decisão de investimento, e sua interpretação exige conhecimento do contexto macroeconômico, setorial e individual de cada empresa.\n\nO investidor é o único responsável por suas decisões financeiras. Antes de realizar qualquer investimento, recomenda-se consultar um profissional devidamente habilitado pela CVM e credenciado pela ANBIMA ou CFP.\n\nEste relatório não substitui a leitura dos documentos oficiais publicados pelas empresas, como Demonstrações Financeiras (DFP), Informes Trimestrais (ITR) e Formulários de Referência (FRE), todos disponíveis no portal da CVM em dados.cvm.gov.br."""

SIGLAS_TERMOS = [
    ("B3",
     "Bolsa de Valores do Brasil. Mercado organizado onde são negociadas ações, fundos e outros ativos financeiros."),
    ("B.E.S.S.T. — Setores de Renda Passiva",
     "Agrupamento setorial composto por Bancos, Energia, Saneamento, Seguros e Telecom. Esses setores são amplamente reconhecidos por características que favorecem a geração de renda passiva: demanda perene independente de ciclos econômicos, concessões de longo prazo, regulação estável e histórico consistente de distribuição de proventos aos acionistas."),
    ("CAGR de Dividendos",
     "Taxa de Crescimento Anual Composta dos dividendos pagos em um período. Indica se os proventos estão crescendo ou diminuindo ao longo do tempo."),
    ("Consistência de Pagamento",
     "Percentual de anos, dentro do período analisado, em que a empresa efetivamente pagou dividendos."),
    ("CVM",
     "Comissão de Valores Mobiliários. Orgão regulador do mercado de capitais brasileiro."),
    ("DY — Dividend Yield",
     "Relação entre os dividendos pagos nos últimos 12 meses e o preço atual da ação. Indica o rendimento em proventos sobre o valor investido."),
    ("Buy and Hold",
     "Estratégia de investimento que consiste em adquirir ações de empresas sólidas e mantê-las por longos períodos, independentemente das oscilações de curto prazo do mercado. O retorno é obtido principalmente pela valorização do patrimônio ao longo do tempo e pelo recebimento contínuo de dividendos, e não pela especulação sobre variações de preço."),
    ("Moat",
     "Termo em inglês para fosso. Representa a vantagem competitiva duradoura de uma empresa, que dificulta a entrada de concorrentes."),
    ("Novo Mercado",
     "Segmento de listagem da B3 com as regras mais rígidas de governança corporativa e maior proteção aos acionistas minoritários."),
    ("P/L — Preco sobre Lucro",
     "Indica quantos anos de lucro atual seriam necessários para recuperar o valor pago pela ação. Valores menores tendem a indicar ação mais barata."),
    ("P/VP — Preco sobre Valor Patrimonial",
     "Compara o preço de mercado da ação com o valor contábil do patrimônio líquido por ação. Valor abaixo de 1 indica negociação com desconto sobre o patrimônio."),
    ("Payout",
     "Percentual do lucro líquido distribuído como dividendos. Payout elevado significa maior distribuição, mas pode limitar reinvestimentos na empresa."),
    ("ROE — Return on Equity",
     "Retorno sobre o Patrimônio Líquido. Mede a eficiência da empresa em gerar lucro a partir do capital dos acionistas."),
    ("Tag Along",
     "Direito do acionista minoritário de vender suas ações nas mesmas condições do acionista controlador em caso de mudança de controle da empresa."),
    ("Web Scraping",
     "Técnica automatizada de coleta de dados de páginas da internet, utilizada neste relatório para obter indicadores de portais financeiros públicos."),
]


# --- SEÇÃO 3: Funções de dados ---

def criar_pastas():
    for p in [PASTA_RESULTADOS, PASTA_RELATORIOS, PASTA_CACHE]:
        os.makedirs(p, exist_ok=True)


def converter_numero_br(valor, eh_percentual=False, tem_decimal_implicito=False):
    if isinstance(valor, (int, float)):
        if isinstance(valor, float) and math.isnan(valor):
            return float("nan")
        return float(valor)
    texto = str(valor).strip()
    if texto in ("", "-", "—", "N/A", "n/a", "nan", "NaN", "000", "0000", "0"):
        return float("nan")
    try:
        if "%" in texto:
            texto = texto.replace("%", "").strip()
            texto = texto.replace(".", "").replace(",", ".")
            return float(texto) / 100
        elif "," in texto:
            partes = texto.split(",")
            inteiro = partes[0].replace(".", "")
            decimal = partes[1] if len(partes) > 1 else "0"
            return float(f"{inteiro}.{decimal}")
        else:
            if tem_decimal_implicito:
                return float(texto) / 100
            return float(texto)
    except (ValueError, IndexError):
        return float("nan")


COLUNAS_DI  = {"p/l","p/vp","psr","p/ativo","p/cap.giro","p/ebit",
               "p/ativ circ.liq","ev/ebit","ev/ebitda","liq.corr",
               "div.liq.patrim","cotacao"}
COLUNAS_PCT = {"dy","roe","roic","mrg.bruta","mrg.ebit","mrg.liq","cresc.rec.5a"}


def buscar_fundamentus():
    print("  Conectando ao Fundamentus.com.br...")
    url = "https://www.fundamentus.com.br/resultado.php"
    cab = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        "Referer": "https://www.fundamentus.com.br/",
    }
    r = requests.get(url, headers=cab, timeout=30)
    r.raise_for_status()
    lista = pd.read_html(StringIO(r.text), flavor="lxml")
    df = lista[0]
    for col in df.columns:
        df[col] = df[col].astype(str)
    df.columns = [str(c).strip().lower() for c in df.columns]
    if "papel" in df.columns:
        df = df.set_index("papel")
    mapa = {
        "div.yield":"dy","p/l":"p/l","p/vp":"p/vp","psr":"psr",
        "p/ativo":"p/ativo","p/cap.giro":"p/cap.giro","p/ebit":"p/ebit",
        "p/ativ circ.liq":"p/ativ circ.liq","ev/ebit":"ev/ebit",
        "ev/ebitda":"ev/ebitda","mrg bruta":"mrg.bruta","mrg ebit":"mrg.ebit",
        "mrg. líq.":"mrg.liq","liq. corr.":"liq.corr","roic":"roic","roe":"roe",
        "liq.2meses":"liq.2meses","patrim. líq":"patrim.liq",
        "dív.líq/ patrim.":"div.liq.patrim","cresc. rec.5a":"cresc.rec.5a",
        "cotação":"cotacao",
    }
    df = df.rename(columns=mapa)
    for col in (COLUNAS_DI | COLUNAS_PCT | {"liq.2meses","patrim.liq"}):
        if col in df.columns:
            df[col] = df[col].apply(
                lambda v: converter_numero_br(v, col in COLUNAS_PCT, col in COLUNAS_DI)
            )
    df = df[df.index.notna() & (df.index.astype(str).str.strip() != "")]
    print(f"  {len(df)} acoes obtidas do Fundamentus")
    return df


def buscar_ibovespa():
    url = (
        "https://sistemaswebb3-listados.b3.com.br/indexProxy/indexCall/"
        "GetPortfolioDay/eyJsYW5ndWFnZSI6InB0LWJyIiwicGFnZU51bWJlciI6MSwi"
        "cGFnZVNpemUiOjEyMCwiaW5kZXgiOiJJQk9WIiwic2VnbWVudCI6IjIifQ=="
    )
    try:
        r = requests.get(url, headers={"User-Agent": "Mozilla/5.0"}, timeout=15)
        r.raise_for_status()
        tickers = [i["cod"] for i in r.json()["results"]]
        print(f"  {len(tickers)} acoes no Ibovespa")
        return tickers
    except Exception as e:
        print(f"  Fallback lista manual ({e})")
        return TODOS_TICKERS_BESST


def aplicar_filtros(df, tickers_ibov):
    df = df[df.index.isin(tickers_ibov)].copy()
    for col in ["liq.2meses","p/l","dy","p/vp"]:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors="coerce")
    if "liq.2meses" in df.columns:
        df = df[df["liq.2meses"] >= FILTRO_LIQUIDEZ_MINIMA]
    if "p/l" in df.columns:
        df = df[(df["p/l"] > FILTRO_PL_MINIMO) & (df["p/l"] <= FILTRO_PL_MAXIMO)]
    if "dy" in df.columns:
        df = df[df["dy"] >= FILTRO_DY_MINIMO]
    if "p/vp" in df.columns:
        df = df[df["p/vp"] <= FILTRO_PVPA_MAXIMO]
    if FILTRO_BESST_ATIVO:
        df = df[df.index.isin(TODOS_TICKERS_BESST)]
    df["setor"] = df.index.map(
        lambda t: next((s for s, l in TICKERS_BESST.items() if t in l), "Outro")
    )
    print(f"  {len(df)} acoes apos filtros")
    return df


def calcular_ranking(df):
    df = df.copy()
    indicadores = {
        "dy":("dy",True),"roe":("roe",True),
        "pl":("p/l",False),"pvpa":("p/vp",False),
        "div_ebitda":("div.liq.patrim",False),
    }
    for met, (col, maior) in indicadores.items():
        if col not in df.columns:
            continue
        s = pd.to_numeric(df[col], errors="coerce").dropna()
        if s.empty or s.max() == s.min():
            df[f"nota_{met}"] = 0.5
            continue
        vmin, vmax = s.min(), s.max()
        df[f"nota_{met}"] = ((df[col]-vmin)/(vmax-vmin)) if maior else ((vmax-df[col])/(vmax-vmin))
    df["nota_final"] = sum(
        df[f"nota_{m}"].fillna(0) * p
        for m, p in PESOS.items() if f"nota_{m}" in df.columns
    )
    df = df.sort_values("nota_final", ascending=False)
    df.insert(0, "ranking", range(1, len(df)+1))
    return df


def buscar_historico_dividendos(tickers):
    resultados = []
    data_ini = datetime.today() - timedelta(days=ANOS_HISTORICO*365)
    for ticker in tickers:
        try:
            divs = yf.Ticker(f"{ticker}.SA").dividends
            if divs.empty:
                resultados.append({"ticker":ticker,"cagr_div":float("nan"),"consistencia_div":0.0,"anos_com_div":0})
                time.sleep(0.3); continue
            divs.index = divs.index.tz_localize(None)
            divs = divs[divs.index >= pd.Timestamp(data_ini)]
            if divs.empty:
                resultados.append({"ticker":ticker,"cagr_div":float("nan"),"consistencia_div":0.0,"anos_com_div":0})
                time.sleep(0.3); continue
            da = divs.groupby(divs.index.year).sum()
            anos_div = len(da)
            consist  = min(anos_div / ANOS_HISTORICO, 1.0)
            cagr = (da.iloc[-1]/da.iloc[0])**(1/(len(da)-1))-1 if len(da)>=2 and da.iloc[0]>0 else float("nan")
            resultados.append({"ticker":ticker,"cagr_div":cagr,"consistencia_div":consist,"anos_com_div":anos_div})
            time.sleep(0.5)
        except Exception:
            resultados.append({"ticker":ticker,"cagr_div":float("nan"),"consistencia_div":0.0,"anos_com_div":0})
            time.sleep(0.5)
    return pd.DataFrame(resultados).set_index("ticker")


# --- SEÇÃO 4: Análise via Gemini (1 chamada por empresa) ---

def inicializar_gemini():
    chave = os.environ.get("GEMINI_API_KEY")
    if not chave:
        print("  GEMINI_API_KEY nao encontrada — secoes qualitativas omitidas")
        return None
    return genai.Client(api_key=chave)


def chamar_gemini(cliente, prompt, max_tent=5):
    """
    Função central de chamada à API Gemini.
    Se trocarmos de modelo no futuro, alteramos só aqui.

    Tratamento de erros:
    - Erro 503 (UNAVAILABLE): modelo sobrecarregado — espera progressiva longa
    - Erro 429 (RESOURCE_EXHAUSTED): quota atingida — espera e tenta novamente
    - Outros erros: até max_tent tentativas com espera crescente
    """
    for t in range(1, max_tent+1):
        try:
            r = cliente.models.generate_content(model=MODELO_GEMINI, contents=prompt)
            return r.text.strip()
        except Exception as e:
            msg = str(e)
            if t < max_tent:
                # Erro 503: modelo sobrecarregado — espera mais longa
                if "503" in msg or "UNAVAILABLE" in msg:
                    espera = 20 * t
                    print(f"      Gemini sobrecarregado (503) — aguardando {espera}s antes de tentar novamente ({t}/{max_tent})...")
                # Erro 429: quota — espera progressiva
                elif "429" in msg or "RESOURCE_EXHAUSTED" in msg:
                    espera = 30 * t
                    print(f"      Quota atingida (429) — aguardando {espera}s ({t}/{max_tent})...")
                else:
                    espera = 10 * t
                    print(f"      Erro na API — aguardando {espera}s ({t}/{max_tent})...")
                time.sleep(espera)
            else:
                return f"[Erro na API apos {max_tent} tentativas: {e}]"


def analisar_empresa(cliente, ticker, setor, dados):
    """
    Faz UMA única chamada ao Gemini por empresa, retornando
    todos os campos em JSON estruturado.
    Isso reduz o consumo do free tier de 3 chamadas para 1 por empresa.
    """
    dy_fmt   = f"{dados.get('dy',0):.1%}"   if pd.notna(dados.get('dy'))   else "N/D"
    pl_fmt   = f"{dados.get('p/l',0):.1f}"  if pd.notna(dados.get('p/l'))  else "N/D"
    pvp_fmt  = f"{dados.get('p/vp',0):.2f}" if pd.notna(dados.get('p/vp')) else "N/D"
    roe_fmt  = f"{dados.get('roe',0):.1%}"  if pd.notna(dados.get('roe'))  else "N/D"
    cagr_val = dados.get('cagr_div', float("nan"))
    cagr_fmt = f"{float(cagr_val):.1%}" if not math.isnan(float(cagr_val)) else "N/D"
    cons_val = dados.get('consistencia_div', 0)
    cons_fmt = f"{float(cons_val):.0%}" if cons_val else "N/D"

    prompt = f"""Responda APENAS com um objeto JSON valido. Sem texto antes ou depois do JSON.
Sem preambulos, sem apresentacoes, sem "Como analista...".
Comece diretamente com o caractere {{ e termine com }}.

Analise a acao {ticker} do setor {setor} para uma carteira previdenciaria
de dividendos de longo prazo baseada na estrategia buy and hold, com foco em
empresas de setores perenes (Bancos, Energia, Saneamento, Seguros e Telecom).

Dados quantitativos:
- DY atual: {dy_fmt} (minimo Barsi: 6%)
- P/L: {pl_fmt}
- P/VP: {pvp_fmt}
- ROE: {roe_fmt}
- CAGR Dividendos 5 anos: {cagr_fmt}
- Consistencia de pagamento: {cons_fmt}

Retorne exatamente este JSON preenchido (todos os campos em portugues do Brasil):

{{
  "modelo_de_negocio": "Como a empresa ganha dinheiro e suas principais fontes de receita. Maximo 3 frases.",
  "vantagem_competitiva": "Qual o diferencial competitivo (moat) desta empresa no setor {setor}. Maximo 3 frases.",
  "perspectiva_do_setor": "Tendencias e perspectivas para o setor {setor} nos proximos 2-3 anos no Brasil. Maximo 3 frases.",
  "principais_riscos": [
    "Risco 1 descrito objetivamente",
    "Risco 2 descrito objetivamente",
    "Risco 3 descrito objetivamente"
  ],
  "segmento_listagem_b3": "Segmento onde a acao esta listada e o que isso significa para o minoritario. 2 frases.",
  "controle_acionario": "Estrutura de controle em 1 frase (estatal, familiar, disperso ou fundo).",
  "historico_gestao": "Avaliacao do historico da gestao em relacao aos acionistas minoritarios. 2 frases.",
  "alertas_recentes": "Escândalos, investigacoes ou noticias negativas dos ultimos 12 meses. Se nao houver, escreva: Nenhum alerta identificado.",
  "nota_governanca": 4,
  "nota_governanca_justificativa": "Justificativa da nota em 1 frase. Escala: 1=pessima ate 5=excelente.",
  "pontos_fortes": [
    "Ponto forte 1 para a tese previdenciaria",
    "Ponto forte 2 para a tese previdenciaria",
    "Ponto forte 3 para a tese previdenciaria"
  ],
  "pontos_atencao": [
    "Ponto de atencao 1",
    "Ponto de atencao 2",
    "Ponto de atencao 3"
  ],
  "adequacao_perfil": "ALTA",
  "adequacao_justificativa": "Justificativa da adequacao ao perfil previdenciario de longo prazo em 2 frases."
}}"""

    texto = chamar_gemini(cliente, prompt)

    # Remove possíveis marcadores de código que o modelo possa inserir
    texto = texto.strip()
    if texto.startswith("```"):
        texto = texto.split("\n", 1)[-1]
    if texto.endswith("```"):
        texto = texto.rsplit("```", 1)[0]

    try:
        return json.loads(texto)
    except json.JSONDecodeError:
        # Se o JSON vier malformado, retorna estrutura de fallback
        return {
            "modelo_de_negocio": texto,
            "vantagem_competitiva": "",
            "perspectiva_do_setor": "",
            "principais_riscos": [],
            "segmento_listagem_b3": "",
            "controle_acionario": "",
            "historico_gestao": "",
            "alertas_recentes": "Nao foi possivel processar a resposta da API.",
            "nota_governanca": 0,
            "nota_governanca_justificativa": "",
            "pontos_fortes": [],
            "pontos_atencao": [],
            "adequacao_perfil": "N/D",
            "adequacao_justificativa": "",
        }


# --- SEÇÃO 5: Formatação Word ---

def cel_cabecalho(celula, texto):
    tc = celula._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "1F497D")
    tcPr.append(shd)
    p = celula.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(texto)
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def cel_dados(celula, texto, negrito=False, centro=True, cor=None):
    p = celula.paragraphs[0]
    if centro:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(str(texto))
    r.font.size = Pt(9)
    r.bold = negrito
    if cor:
        r.font.color.rgb = cor


def zebra(linha, par):
    if par % 2 == 0:
        for c in linha.cells:
            tc = c._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "F2F2F2")
            tcPr.append(shd)


def avaliar_checklist(dados):
    """Avalia cada critério do checklist Barsi e retorna lista de tuplas."""
    itens = []
    def sf(v):
        try: return float(v)
        except: return float("nan")

    dy   = sf(dados.get("dy"))
    pl   = sf(dados.get("p/l"))
    pvp  = sf(dados.get("p/vp"))
    roe  = sf(dados.get("roe"))
    cagr = sf(dados.get("cagr_div"))
    cons = sf(dados.get("consistencia_div", 0))
    anos = int(sf(dados.get("anos_com_div", 0)))
    dvp  = sf(dados.get("div.liq.patrim"))

    def ico(ok, warn=False):
        return "OK" if ok else ("ATT" if warn else "NAO")
    def st(ok, warn=False):
        return "ok" if ok else ("alerta" if warn else "risco")

    if not math.isnan(dvp):
        ok = dvp < 2.5
        itens.append((ico(ok), "Divida Liq./Patrim. < 2,5x", f"{dvp:.2f}x", st(ok)))
    else:
        itens.append(("ATT", "Divida Liq./Patrim. < 2,5x", "N/D", "alerta"))

    if not math.isnan(roe):
        ok = roe > 0.12
        itens.append((ico(ok), "ROE > 12%", f"{roe:.1%}", st(ok)))

    itens.append((
        ico(anos >= 5, anos >= 3),
        "Historico > 5 anos consecutivos",
        f"{anos} anos",
        st(anos >= 5, anos >= 3)
    ))

    if not math.isnan(dy):
        ok = dy >= 0.06
        itens.append((ico(ok), "DY >= 6% (criterio minimo previdenciario)", f"{dy:.1%}", st(ok)))

    if not math.isnan(cagr):
        ok = cagr > 0
        itens.append((ico(ok, True), "Dividendo crescendo (CAGR > 0)", f"{cagr:.1%} a.a.", st(ok, True)))

    if not math.isnan(pl):
        ok = 5 <= pl <= 20
        itens.append((ico(ok, True), "P/L em faixa razoavel (5-20x)", f"{pl:.1f}x", st(ok, True)))

    if not math.isnan(pvp):
        ok = pvp <= 2.0
        itens.append((ico(ok, True), "P/VP <= 2,0x", f"{pvp:.2f}x", st(ok, True)))

    return itens


def add_paragrafo(doc, texto, tamanho=10, negrito=False, cor=None,
                  espaco_antes=0, espaco_depois=4, recuo=0, italico=False,
                  justificado=False):
    """
    Cria um parágrafo formatado no documento Word.
    O parâmetro justificado=True aplica alinhamento justificado,
    usado nos textos descritivos das empresas para aparência profissional.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(espaco_antes)
    p.paragraph_format.space_after  = Pt(espaco_depois)
    if recuo:
        p.paragraph_format.left_indent = Cm(recuo)
    if justificado:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(texto)
    r.font.size = Pt(tamanho)
    r.bold    = negrito
    r.italic  = italico
    if cor:
        r.font.color.rgb = cor
    return p


# --- SEÇÃO 6: Sumário dinâmico ---

def calcular_pagina_estimada(pos_empresa):
    """
    Estima o número de página de cada empresa no documento.
    Estrutura fixa:
      Pág 1: Capa
      Pág 2: Sumário
      Pág 3: Introdução
      Pág 4: Ranking TOP 10
      Pág 5+: Análise TOP 5 (aprox. 3 páginas por empresa)
    """
    pagina_inicio_analise = 5
    paginas_por_empresa   = 3
    return pagina_inicio_analise + (pos_empresa - 1) * paginas_por_empresa


def gerar_sumario(doc, empresas_top5):
    """
    Gera o sumário dinamicamente com os tickers e setores reais
    do ranking daquela semana.

    Parâmetros:
        doc         : documento Word em construção
        empresas_top5: lista de tuplas (ticker, setor) das TOP 5
    """
    h = doc.add_heading("SUMARIO", level=1)
    h.runs[0].font.color.rgb = COR_TITULO

    # Estrutura fixa do documento
    itens_fixos = [
        ("Introducao e Objetivos do Relatorio", 3),
        ("Secao 1 — Ranking TOP 10 — Setores de Renda Passiva", 4),
        ("Secao 2 — Analise Aprofundada TOP 5",  5),
    ]
    for titulo, pag in itens_fixos:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r_titulo = p.add_run(titulo)
        r_titulo.font.size = Pt(10)
        # Tab stop para alinhar número de página à direita
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15))
        r_pag = p.add_run(f"\t{pag}")
        r_pag.font.size = Pt(10)

    # Subseções dinâmicas — um item por empresa do TOP 5
    for i, (ticker, setor) in enumerate(empresas_top5, 1):
        pag_estimada = calcular_pagina_estimada(i)
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.left_indent  = Cm(1)
        r_titulo = p.add_run(f"  2.{i}  {ticker}  —  {setor}")
        r_titulo.font.size = Pt(10)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15))
        r_pag = p.add_run(f"\t{pag_estimada}")
        r_pag.font.size = Pt(10)

    # Itens finais fixos
    pag_siglas  = 5 + len(empresas_top5) * 3
    pag_metad   = pag_siglas + 1

    for titulo, pag in [
        ("Siglas e Termos", pag_siglas),
        ("Metadados do Relatorio", pag_metad),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r_titulo = p.add_run(titulo)
        r_titulo.font.size = Pt(10)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15))
        r_pag = p.add_run(f"\t{pag}")
        r_pag.font.size = Pt(10)

    nota = doc.add_paragraph()
    nota.paragraph_format.space_before = Pt(8)
    rn = nota.add_run(
        "Nota: os numeros de pagina sao estimativas. Ao abrir o arquivo, "
        "clique com o botao direito no sumario e selecione 'Atualizar campo' "
        "para obter a paginacao exata."
    )
    rn.font.size  = Pt(8)
    rn.italic     = True
    rn.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


# --- SEÇÃO 7: Geração do documento ---

def gerar_relatorio(df_ranking, df_historico, cliente_gemini):
    doc   = Document()
    agora = datetime.now()

    # Página A4
    sec = doc.sections[0]
    sec.page_width    = Cm(21.0)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)

    # Mescla ranking com histórico
    df_top = df_ranking.head(TOP_RANKING).copy()
    df_top = df_top.merge(df_historico, left_index=True, right_index=True, how="left")

    # Lista dinâmica das TOP 5 para o sumário
    df_top5     = df_top.head(TOP_ANALISE)
    empresas_top5 = [(row.name, row.get("setor","")) for _, row in df_top5.iterrows()]

    # ── CAPA ─────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    r = p.add_run("RELATORIO SEMANAL DE INVESTIMENTOS")
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = COR_TITULO

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Carteira Previdenciaria — Estrategia de Dividendos de Longo Prazo")
    r2.font.size = Pt(14); r2.font.color.rgb = COR_SUBTITULO

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(f"Semana de {agora.strftime('%d/%m/%Y')}")
    r3.font.size = Pt(11)

    doc.add_page_break()

    # ── SUMÁRIO DINÂMICO ─────────────────────────────────────────────────
    gerar_sumario(doc, empresas_top5)
    doc.add_page_break()

    # ── INTRODUÇÃO ───────────────────────────────────────────────────────
    h_intro = doc.add_heading("INTRODUCAO E OBJETIVOS DO RELATORIO", level=1)
    h_intro.runs[0].font.color.rgb = COR_TITULO

    add_paragrafo(doc, TEXTO_INTRODUCAO, espaco_depois=8)

    secoes_intro = [
        ("Objetivo",                                          TEXTO_OBJETIVO),
        ("Metodologia",                                       TEXTO_METODOLOGIA),
        ("Limitacoes e Defasagem de Dados",                   TEXTO_LIMITACOES),
        ("Aviso Importante — Nao Constitui Recomendacao de Investimento", TEXTO_AVISO),
    ]
    for titulo_sec, texto_sec in secoes_intro:
        add_paragrafo(doc, titulo_sec, tamanho=11, negrito=True,
                      cor=COR_SUBTITULO, espaco_antes=8, espaco_depois=4)
        for bloco in texto_sec.split("\n\n"):
            if bloco.strip():
                add_paragrafo(doc, bloco.strip(), espaco_depois=6)

    doc.add_page_break()

    # ── SEÇÃO 1: RANKING TOP 10 ──────────────────────────────────────────
    h1 = doc.add_heading("SECAO 1 — Ranking TOP 10 — Setores de Renda Passiva", level=1)
    h1.runs[0].font.color.rgb = COR_TITULO

    p_param = doc.add_paragraph()
    r_label = p_param.add_run("Filtros aplicados: ")
    r_label.bold = True; r_label.font.size = Pt(9)
    r_val = p_param.add_run(
        f"DY >= {FILTRO_DY_MINIMO:.0%}  |  "
        f"P/L entre {FILTRO_PL_MINIMO} e {FILTRO_PL_MAXIMO}  |  "
        f"P/VP <= {FILTRO_PVPA_MAXIMO}  |  "
        f"Liquidez >= R$ {FILTRO_LIQUIDEZ_MINIMA:,.0f}  |  "
        f"Setores: {'Renda Passiva (B, E, S, S, T)' if FILTRO_BESST_ATIVO else 'Todos'}"
    )
    r_val.font.size = Pt(9); r_val.font.color.rgb = RGBColor(0x60,0x60,0x60)

    doc.add_paragraph()

    # Tabela ranking
    cols_rank = ["#","Ticker","Setor","DY","P/L","P/VP","ROE","CAGR Div","Consist.","Nota"]
    tabela = doc.add_table(rows=1, cols=len(cols_rank))
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, c in enumerate(cols_rank):
        cel_cabecalho(tabela.rows[0].cells[i], c)

    for _, row in df_top.iterrows():
        ticker = row.name
        cagr_v = row.get("cagr_div", float("nan"))
        cons_v = row.get("consistencia_div", 0)
        cagr_s = f"{float(cagr_v):.1%}" if not math.isnan(float(cagr_v)) else "N/D"
        cons_s = f"{float(cons_v):.0%}"  if cons_v else "N/D"
        dy_v   = row.get("dy")
        rank_n = int(row["ranking"])

        nova = tabela.add_row()
        vals = [
            rank_n,
            ticker,
            row.get("setor",""),
            f"{dy_v:.1%}"         if pd.notna(dy_v)              else "N/D",
            f"{row['p/l']:.1f}"   if pd.notna(row.get("p/l"))    else "N/D",
            f"{row['p/vp']:.2f}"  if pd.notna(row.get("p/vp"))   else "N/D",
            f"{row['roe']:.1%}"   if pd.notna(row.get("roe"))     else "N/D",
            cagr_s, cons_s,
            f"{row['nota_final']:.3f}",
        ]
        for i, v in enumerate(vals):
            cor = None
            if i == 3 and pd.notna(dy_v):
                cor = COR_DESTAQUE if dy_v >= 0.08 else RGBColor(0xBF,0x8F,0x00)
            cel_dados(nova.cells[i], v, negrito=(i==1), centro=(i!=2), cor=cor)
        zebra(nova, rank_n)

    add_paragrafo(doc,
        "Metodologia: ranking multicriterio com normalizacao min-max. "
        "Pesos: DY 30% | ROE 25% | P/L 20% | P/VP 15% | Div/Patrim 10%. "
        "Nota varia de 0 (pior) a 1 (melhor).",
        tamanho=8, italico=True, cor=RGBColor(0x80,0x80,0x80), espaco_antes=6)

    doc.add_page_break()

    # ── SEÇÃO 2: ANÁLISE TOP 5 ───────────────────────────────────────────
    h2_sec = doc.add_heading("SECAO 2 — Analise Aprofundada TOP 5", level=1)
    h2_sec.runs[0].font.color.rgb = COR_TITULO
    add_paragrafo(doc,
        "Analise detalhada das 5 melhores acoes do ranking, combinando "
        "indicadores quantitativos e analise qualitativa via inteligencia artificial.",
        tamanho=10)

    for pos, (_, row) in enumerate(df_top5.iterrows(), 1):
        ticker = row.name
        setor  = row.get("setor", "")
        dados  = {
            "dy": row.get("dy"), "p/l": row.get("p/l"), "p/vp": row.get("p/vp"),
            "roe": row.get("roe"), "div.liq.patrim": row.get("div.liq.patrim"),
            "cagr_div": row.get("cagr_div"), "consistencia_div": row.get("consistencia_div"),
            "anos_com_div": row.get("anos_com_div", 0),
        }

        doc.add_page_break()

        h_emp = doc.add_heading(f"2.{pos}  {ticker}  —  {setor}", level=2)
        h_emp.runs[0].font.color.rgb = COR_SUBTITULO

        # Tabela de identificação
        tab_id = doc.add_table(rows=2, cols=4)
        tab_id.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, c in enumerate(["Ticker","Setor","DY Atual","Nota Ranking"]):
            cel_cabecalho(tab_id.rows[0].cells[i], c)
        for i, v in enumerate([
            ticker, setor,
            f"{row['dy']:.1%}" if pd.notna(row.get("dy")) else "N/D",
            f"{row['nota_final']:.3f}"
        ]):
            cel_dados(tab_id.rows[1].cells[i], v, negrito=True)

        doc.add_paragraph()

        # Chamada única ao Gemini
        analise = {}
        if cliente_gemini:
            print(f"    [{pos}/5] Chamando Gemini para {ticker}...")
            analise = analisar_empresa(cliente_gemini, ticker, setor, dados)
            time.sleep(5)  # Pausa entre empresas — respeita rate limit do free tier
        else:
            analise = {k: "[API nao configurada]" for k in
                       ["modelo_de_negocio","vantagem_competitiva","perspectiva_do_setor",
                        "segmento_listagem_b3","controle_acionario","historico_gestao",
                        "alertas_recentes","nota_governanca_justificativa",
                        "adequacao_besst","adequacao_justificativa"]}
            analise["principais_riscos"]  = []
            analise["pontos_fortes"]      = []
            analise["pontos_atencao"]     = []
            analise["nota_governanca"]    = 0

        # 2.x.1 O Negócio
        add_paragrafo(doc, "O Negocio", tamanho=11, negrito=True,
                      cor=COR_SUBTITULO, espaco_antes=8, espaco_depois=2)

        for campo, rotulo in [
            ("modelo_de_negocio",   "Modelo de Negocio"),
            ("vantagem_competitiva","Vantagem Competitiva"),
            ("perspectiva_do_setor","Perspectiva do Setor"),
        ]:
            add_paragrafo(doc, rotulo, tamanho=10, negrito=True, espaco_depois=1)
            add_paragrafo(doc, analise.get(campo, ""), tamanho=10,
                          recuo=0.5, espaco_depois=4, justificado=True)

        add_paragrafo(doc, "Principais Riscos", tamanho=10, negrito=True, espaco_depois=1)
        for risco in analise.get("principais_riscos", []):
            add_paragrafo(doc, f"- {risco}", tamanho=10, recuo=0.5,
                          espaco_depois=2, justificado=True)

        doc.add_paragraph()

        # 2.x.2 Saúde Financeira
        add_paragrafo(doc, "Saude Financeira", tamanho=11, negrito=True,
                      cor=COR_SUBTITULO, espaco_antes=4, espaco_depois=2)

        tab_fin = doc.add_table(rows=5, cols=3)
        tab_fin.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, c in enumerate(["Indicador","Valor","Status"]):
            cel_cabecalho(tab_fin.rows[0].cells[i], c)

        def si(cond): return "Adequado" if cond else "Atenção"

        fin = [
            ("DY Atual",  f"{row['dy']:.1%}"   if pd.notna(row.get('dy'))  else "N/D",
             si(pd.notna(row.get('dy'))  and row['dy']  >= 0.06)),
            ("P/L",       f"{row['p/l']:.1f}"  if pd.notna(row.get('p/l')) else "N/D",
             si(pd.notna(row.get('p/l')) and 5 <= row['p/l'] <= 20)),
            ("ROE",       f"{row['roe']:.1%}"  if pd.notna(row.get('roe')) else "N/D",
             si(pd.notna(row.get('roe')) and row['roe']  > 0.12)),
            ("P/VP",      f"{row['p/vp']:.2f}" if pd.notna(row.get('p/vp'))else "N/D",
             si(pd.notna(row.get('p/vp'))and row['p/vp'] <= 2.0)),
        ]
        for i, (ind, val, st) in enumerate(fin, 1):
            cel_dados(tab_fin.rows[i].cells[0], ind)
            cel_dados(tab_fin.rows[i].cells[1], val, negrito=True)
            cel_dados(tab_fin.rows[i].cells[2], st)

        doc.add_paragraph()

        # 2.x.3 Histórico de Dividendos
        add_paragrafo(doc, "Historico de Dividendos", tamanho=11, negrito=True,
                      cor=COR_SUBTITULO, espaco_antes=4, espaco_depois=2)

        cagr_v = row.get("cagr_div", float("nan"))
        cons_v = row.get("consistencia_div", 0)
        anos_v = int(row.get("anos_com_div", 0))
        cagr_s = f"{float(cagr_v):.1%}" if not math.isnan(float(cagr_v)) else "N/D"
        cons_s = f"{float(cons_v):.0%}"  if cons_v else "N/D"

        tab_div = doc.add_table(rows=4, cols=3)
        tab_div.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, c in enumerate(["Indicador","Valor","Status"]):
            cel_cabecalho(tab_div.rows[0].cells[i], c)
        div_rows = [
            ("CAGR Dividendos 5a",     cagr_s, si(not math.isnan(float(cagr_v)) and float(cagr_v)>0)),
            ("Consistencia Pagamento", cons_s, si(float(cons_v) >= 0.8)),
            (f"Anos com Dividendos",   f"{anos_v}/{ANOS_HISTORICO}", si(anos_v >= 4)),
        ]
        for i, (ind, val, st) in enumerate(div_rows, 1):
            cel_dados(tab_div.rows[i].cells[0], ind)
            cel_dados(tab_div.rows[i].cells[1], val, negrito=True)
            cel_dados(tab_div.rows[i].cells[2], st)

        doc.add_paragraph()

        # 2.x.4 Checklist Barsi
        add_paragrafo(doc, "Checklist Barsi", tamanho=11, negrito=True,
                      cor=COR_SUBTITULO, espaco_antes=4, espaco_depois=2)

        for ico, texto_c, val_c, status_c in avaliar_checklist(dados):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_after = Pt(2)
            r_ico = p.add_run(f"[{ico}]  ")
            r_ico.bold = True
            r_ico.font.size = Pt(10)
            if status_c == "ok":
                r_ico.font.color.rgb = COR_DESTAQUE
            elif status_c == "alerta":
                r_ico.font.color.rgb = RGBColor(0xBF, 0x8F, 0x00)
            else:
                r_ico.font.color.rgb = COR_RISCO
            r_txt = p.add_run(texto_c)
            r_txt.font.size = Pt(10)
            r_val = p.add_run(f"   {val_c}")
            r_val.font.size = Pt(10)
            r_val.bold = True

        doc.add_paragraph()

        # 2.x.5 Governança
        add_paragrafo(doc, "Governanca Corporativa", tamanho=11, negrito=True,
                      cor=COR_SUBTITULO, espaco_antes=4, espaco_depois=2)

        for campo, rotulo in [
            ("segmento_listagem_b3", "Segmento de Listagem B3"),
            ("controle_acionario",   "Controle Acionario"),
            ("historico_gestao",     "Historico de Gestao"),
            ("alertas_recentes",     "Alertas Recentes"),
        ]:
            add_paragrafo(doc, rotulo, tamanho=10, negrito=True, espaco_depois=1)
            add_paragrafo(doc, analise.get(campo, ""), tamanho=10,
                          recuo=0.5, espaco_depois=4, justificado=True)

        nota_gov = analise.get("nota_governanca", 0)
        nota_just = analise.get("nota_governanca_justificativa", "")
        add_paragrafo(doc, f"Nota de Governanca: {nota_gov}/5", tamanho=10,
                      negrito=True, espaco_depois=1)
        add_paragrafo(doc, nota_just, tamanho=10, recuo=0.5, espaco_depois=4)

        doc.add_paragraph()

        # 2.x.6 Veredito
        add_paragrafo(doc, "Veredito Consolidado", tamanho=11, negrito=True,
                      cor=COR_SUBTITULO, espaco_antes=4, espaco_depois=2)

        adequacao = analise.get("adequacao_perfil", analise.get("adequacao_besst", "N/D"))
        cor_adeq  = (COR_DESTAQUE if adequacao == "ALTA"
                     else RGBColor(0xBF,0x8F,0x00) if adequacao == "MEDIA"
                     else COR_RISCO)

        add_paragrafo(doc, f"Adequacao ao perfil previdenciario: {adequacao}",
                      tamanho=11, negrito=True, cor=cor_adeq, espaco_depois=2)
        add_paragrafo(doc, analise.get("adequacao_justificativa",""),
                      tamanho=10, recuo=0.5, espaco_depois=6, justificado=True)

        add_paragrafo(doc, "Pontos Fortes da Tese Previdenciaria",
                      tamanho=10, negrito=True, espaco_depois=1)
        for pf in analise.get("pontos_fortes", []):
            add_paragrafo(doc, f"- {pf}", tamanho=10, recuo=0.5,
                          espaco_depois=2, justificado=True)

        add_paragrafo(doc, "Pontos de Atencao", tamanho=10, negrito=True,
                      espaco_antes=4, espaco_depois=1)
        for pa in analise.get("pontos_atencao", []):
            add_paragrafo(doc, f"- {pa}", tamanho=10, recuo=0.5,
                          espaco_depois=2, justificado=True)

    # ── SIGLAS E TERMOS ──────────────────────────────────────────────────
    doc.add_page_break()
    h_sig = doc.add_heading("SIGLAS E TERMOS UTILIZADOS NESTE RELATORIO", level=1)
    h_sig.runs[0].font.color.rgb = COR_TITULO

    for termo, definicao in SIGLAS_TERMOS:
        add_paragrafo(doc, termo, tamanho=10, negrito=True,
                      espaco_antes=6, espaco_depois=1)
        add_paragrafo(doc, definicao, tamanho=10, recuo=0.5, espaco_depois=2)

    # ── METADADOS ────────────────────────────────────────────────────────
    doc.add_page_break()
    h_meta = doc.add_heading("Metadados do Relatorio", level=2)
    h_meta.runs[0].font.color.rgb = COR_TITULO

    for k, v in [
        ("Data de geracao",             agora.strftime("%d/%m/%Y %H:%M")),
        ("Fontes de dados",             "Fundamentus.com.br, B3, yfinance, API Gemini"),
        ("Modelo de IA",                MODELO_GEMINI),
        ("Total acoes Ibovespa",        "~79"),
        ("Acoes apos filtros",          str(len(df_ranking))),
        ("Filtro DY minimo",            f"{FILTRO_DY_MINIMO:.0%}"),
        ("Filtro P/L maximo",           str(FILTRO_PL_MAXIMO)),
        ("Filtro setorial (B,E,S,S,T)", "Ativo" if FILTRO_BESST_ATIVO else "Inativo"),
        ("Chamadas a API por relatorio", f"{TOP_ANALISE} (1 por empresa)"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        rk = p.add_run(f"{k}: ")
        rk.bold = True; rk.font.size = Pt(9)
        rv = p.add_run(v)
        rv.font.size = Pt(9)

    add_paragrafo(doc,
        "Este relatorio foi gerado automaticamente para fins educacionais e de apoio "
        "a decisao pessoal de investimento. Nenhuma informacao constitui recomendacao "
        "formal de compra ou venda de valores mobiliarios.",
        tamanho=8, italico=True, cor=RGBColor(0x80,0x80,0x80), espaco_antes=8)

    # ── Salvar ───────────────────────────────────────────────────────────
    nome    = f"relatorio_besst_{agora.strftime('%Y%m%d_%H%M')}.docx"
    caminho = os.path.join(PASTA_RELATORIOS, nome)
    doc.save(caminho)
    return caminho


# --- SEÇÃO 8: Execução principal ---

if __name__ == "__main__":
    inicio = datetime.now()

    print("\n" + "="*60)
    print("  RELATORIO SEMANAL — CARTEIRA DE DIVIDENDOS")
    print("  Projeto B3 — Analise de Investimentos")
    print("="*60)

    criar_pastas()

    print("\nInicializando API Gemini...")
    cliente_gemini = inicializar_gemini()
    if cliente_gemini:
        print("  API Gemini pronta")

    print("\nColetando dados de mercado...")
    tickers_ibov = buscar_ibovespa()
    time.sleep(1)
    df_todos = buscar_fundamentus()

    print("\nAplicando filtros e calculando ranking...")
    df_filtrado = aplicar_filtros(df_todos, tickers_ibov)
    if df_filtrado.empty:
        print("Nenhuma acao passou pelos filtros. Encerrando.")
        exit(1)
    df_ranking = calcular_ranking(df_filtrado)

    print(f"\nBuscando historico de dividendos ({ANOS_HISTORICO} anos)...")
    tickers_hist = df_ranking.head(TOP_ANALISE).index.tolist()
    df_historico = buscar_historico_dividendos(tickers_hist)

    print(f"\nGerando relatorio Word...")
    print(f"  ({TOP_ANALISE} chamadas ao Gemini — 1 por empresa)\n")
    caminho = gerar_relatorio(df_ranking, df_historico, cliente_gemini)

    fim   = datetime.now()
    tempo = (fim - inicio).seconds
    print("\n" + "="*60)
    print(f"  Relatorio gerado em {tempo} segundos!")
    print(f"  {os.path.abspath(caminho)}")
    print("="*60 + "\n")
