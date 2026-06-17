# ============================================================
# NOME DO SCRIPT: relatorio_mensal_besst.py
# OBJETIVO: Gerar relatório mensal em Word (.docx) com:
#           - Ranking TOP 10 com gráfico P/L e P/VP
#           - Gráfico de evolução do DY histórico TOP 5
#           - Tabela de Oportunidades Ocultas TOP 5
# BIBLIOTECAS: python-docx, yfinance, requests, pandas,
#              matplotlib
# EXECUTAR:    python relatorio_mensal_besst.py
# INSTALAR:    pip install python-docx yfinance requests
#              pandas matplotlib
# ============================================================


# --- SEÇÃO 1: Importações ---
import os, time, warnings, math
from datetime import datetime, timedelta
from io import StringIO, BytesIO

import pandas as pd
import requests
import yfinance as yf
import matplotlib
matplotlib.use("Agg")   # Backend sem interface gráfica — funciona no Colab e Windows
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

warnings.filterwarnings("ignore")


# --- SEÇÃO 2: Configurações ---
# ⚙️ Altere esses valores conforme sua análise

PASTA_RESULTADOS  = "resultados"
PASTA_RELATORIOS  = "resultados/relatorios"
PASTA_CACHE       = "dados/fundamentos"

# Filtros do ranking principal (TOP 10)
FILTRO_LIQUIDEZ_MINIMA = 500_000
FILTRO_PL_MINIMO       = 0
FILTRO_PL_MAXIMO       = 40
FILTRO_DY_MINIMO       = 0.06      # 6% — critério mínimo previdenciário
FILTRO_PVPA_MAXIMO     = 5.0
FILTRO_BESST_ATIVO     = True

# Filtros das Oportunidades Ocultas
# Objetivo: empresas BESST desvalorizadas onde é possível acumular
# volume com baixo desembolso por ação — estratégia de acumulação
# de longo prazo. Filtros propositalmente mais flexíveis.
OO_ROE_MINIMO   = 0.08   # ROE > 8% — empresa ainda operacionalmente sadia
OO_PVPA_MAXIMO  = 1.5    # P/VP <= 1,5x — descontada ou próxima do patrimônio
OO_PL_MINIMO    = 0      # Sem prejuízo (P/L deve ser positivo)
OO_TOP          = 5      # Quantas oportunidades exibir
# Sem filtro de liquidez — estratégia de acumulação e manutenção
# Sem filtro de dívida — universo BESST pequeno, restringiria demais

# Tickers por setor B.E.S.S.T.
TICKERS_BESST = {
    "Bancos":     ["BBAS3","BBDC3","BBDC4","ITUB4","ITSA4","SANB11","BPAC11"],
    "Energia":    ["ELET3","ELET6","CMIG4","CPFE3","CPLE6","EGIE3","ENEV3",
                   "ENGI11","EQTL3","TAEE11","ISAE4"],
    "Saneamento": ["SBSP3"],
    "Seguros":    ["BBSE3","IRBR3","PSSA3"],
    "Telecom":    ["VIVT3","TIMS3"],
}
TODOS_TICKERS_BESST = [t for lista in TICKERS_BESST.values() for t in lista]

# Pesos do ranking multicritério
PESOS = {
    "dy": 0.30, "roe": 0.25, "pl": 0.20, "pvpa": 0.15, "div_ebitda": 0.10,
}

# Cores do documento
COR_TITULO    = RGBColor(0x1F, 0x49, 0x7D)
COR_SUBTITULO = RGBColor(0x2E, 0x75, 0xB6)
COR_DESTAQUE  = RGBColor(0x70, 0xAD, 0x47)
COR_RISCO     = RGBColor(0xFF, 0x00, 0x00)

# Cores dos gráficos (paleta harmônica com o documento)
CORES_GRAFICOS = ["#1F497D", "#2E75B6", "#70AD47", "#ED7D31", "#A9D18E",
                  "#9DC3E6", "#FFC000", "#FF0000", "#00B0F0", "#7030A0"]

TOP_RANKING    = 10
TOP_HISTORICO  = 5     # Empresas no gráfico de DY histórico
ANOS_HISTORICO = 5

# Textos da introdução
TEXTO_INTRODUCAO = (
    "Este relatório foi gerado automaticamente por meio de coleta, processamento "
    "e consolidação de dados públicos disponíveis em portais especializados do "
    "mercado financeiro brasileiro, como Fundamentus, Yahoo Finance e a base de "
    "dados oficial da CVM. Nenhuma informação utilizada nesta análise é proprietária "
    "ou restrita — todas as fontes consultadas são de acesso público e gratuito."
)
TEXTO_OBJETIVO = (
    "O propósito deste relatório mensal é reunir, em um único documento estruturado, "
    "os principais indicadores fundamentalistas de empresas listadas na B3, facilitando "
    "a leitura e a comparação entre ativos para o investidor pessoa física.\n\n"
    "A filosofia que norteia a seleção dos ativos é a de acumulação de longo prazo: "
    "adquirir participações em empresas sólidas, com histórico consistente de geração "
    "de caixa e distribuição de proventos, e mantê-las pelo maior prazo possível. "
    "Essa abordagem — conhecida como buy and hold — difere das estratégias especulativas "
    "de curto prazo, que buscam lucro na oscilação de preços. O investidor de longo prazo "
    "constrói renda passiva real e crescente ao longo dos anos, tornando-se progressivamente "
    "menos dependente do trabalho ativo para sustentar seu padrão de vida."
)
TEXTO_METODOLOGIA = (
    "Os dados foram coletados de forma automatizada via Python, utilizando técnicas de "
    "raspagem de dados em portais financeiros públicos. Os indicadores foram calculados "
    "ou extraídos diretamente das fontes originais, sem alteração de conteúdo."
)
TEXTO_LIMITACOES = (
    "Por se tratar de coleta automatizada, os dados podem apresentar defasagem em relação "
    "à publicação oficial mais recente. Recomenda-se sempre verificar os dados diretamente "
    "nas fontes originais antes de qualquer utilização."
)
TEXTO_AVISO = (
    "As informações contidas neste relatório têm finalidade exclusivamente informativa e "
    "educacional. Nenhum conteúdo aqui apresentado representa recomendação de compra, venda "
    "ou manutenção de valores mobiliários. O investidor é o único responsável por suas "
    "decisões financeiras. Antes de realizar qualquer investimento, recomenda-se consultar "
    "um profissional habilitado pela CVM e credenciado pela ANBIMA ou CFP."
)

SIGLAS_TERMOS = [
    ("B3", "Bolsa de Valores do Brasil."),
    ("B.E.S.S.T.", "Agrupamento setorial de Bancos, Energia, Saneamento, Seguros e Telecom. "
     "Setores com demanda perene e histórico consistente de distribuição de proventos."),
    ("Buy and Hold", "Estratégia de adquirir ações e mantê-las por longos períodos, "
     "obtendo retorno pela valorização e pelo recebimento contínuo de dividendos."),
    ("CAGR de Dividendos", "Taxa de Crescimento Anual Composta dos dividendos pagos. "
     "Indica se os proventos estão crescendo ou diminuindo ao longo do tempo."),
    ("CVM", "Comissão de Valores Mobiliários. Orgão regulador do mercado de capitais brasileiro."),
    ("DY — Dividend Yield", "Relação entre dividendos pagos nos últimos 12 meses e o preço "
     "atual da ação. Indica o rendimento em proventos sobre o valor investido."),
    ("Oportunidade Oculta", "Empresa com fundamentos sólidos (ROE elevado, dividendos "
     "crescentes, baixo endividamento) que ainda não foi plenamente reconhecida pelo mercado, "
     "apresentando valuation atrativo."),
    ("P/L — Preco sobre Lucro", "Quantos anos de lucro atual seriam necessários para "
     "recuperar o valor pago pela ação. Valores menores indicam ação mais barata."),
    ("P/VP — Preco sobre Valor Patrimonial", "Compara o preço de mercado com o valor "
     "contábil do patrimônio por ação. Abaixo de 1 indica desconto sobre o patrimônio."),
    ("ROE — Return on Equity", "Retorno sobre o Patrimônio Líquido. Mede a eficiência "
     "da empresa em gerar lucro a partir do capital dos acionistas."),
]


# --- SEÇÃO 3: Funções de dados ---

def criar_pastas():
    """Cria as pastas de saída se não existirem."""
    for p in [PASTA_RESULTADOS, PASTA_RELATORIOS, PASTA_CACHE]:
        os.makedirs(p, exist_ok=True)


def converter_numero_br(valor, eh_percentual=False, tem_decimal_implicito=False):
    """
    Converte valores numéricos no formato do Fundamentus para float.
    O site retorna números como strings no padrão brasileiro.
    """
    if isinstance(valor, (int, float)):
        if isinstance(valor, float) and math.isnan(valor):
            return float("nan")
        return float(valor)
    texto = str(valor).strip()
    if texto in ("", "-", "nan", "NaN", "000", "0000", "0"):
        return float("nan")
    try:
        if "%" in texto:
            texto = texto.replace("%", "").strip().replace(".", "").replace(",", ".")
            return float(texto) / 100
        elif "," in texto:
            partes = texto.split(",")
            return float(f"{partes[0].replace('.','')}.{partes[1] if len(partes)>1 else '0'}")
        else:
            return float(texto) / 100 if tem_decimal_implicito else float(texto)
    except (ValueError, IndexError):
        return float("nan")


COLUNAS_DI  = {"p/l","p/vp","psr","p/ativo","p/cap.giro","p/ebit",
               "p/ativ circ.liq","ev/ebit","ev/ebitda","liq.corr",
               "div.liq.patrim","cotacao"}
COLUNAS_PCT = {"dy","roe","roic","mrg.bruta","mrg.ebit","mrg.liq","cresc.rec.5a"}


def buscar_fundamentus():
    """Busca indicadores do Fundamentus via scraping direto."""
    print("  Conectando ao Fundamentus.com.br...")
    url = "https://www.fundamentus.com.br/resultado.php"
    cab = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        "Referer": "https://www.fundamentus.com.br/",
    }
    r = requests.get(url, headers=cab, timeout=30)
    r.raise_for_status()
    lista = pd.read_html(StringIO(r.text), flavor="lxml")
    df = lista[0]
    for col in df.columns:
        df[col] = df[col].astype(str)
    df.columns = [str(c).strip().lower() for c in df.columns]
    if "papel" in df.columns:
        df = df.set_index("papel")
    mapa = {
        "div.yield":"dy","p/l":"p/l","p/vp":"p/vp","psr":"psr",
        "p/ativo":"p/ativo","p/cap.giro":"p/cap.giro","p/ebit":"p/ebit",
        "p/ativ circ.liq":"p/ativ circ.liq","ev/ebit":"ev/ebit",
        "ev/ebitda":"ev/ebitda","mrg bruta":"mrg.bruta","mrg ebit":"mrg.ebit",
        "mrg. líq.":"mrg.liq","liq. corr.":"liq.corr","roic":"roic","roe":"roe",
        "liq.2meses":"liq.2meses","patrim. líq":"patrim.liq",
        "dív.líq/ patrim.":"div.liq.patrim","cresc. rec.5a":"cresc.rec.5a",
        "cotação":"cotacao",
    }
    df = df.rename(columns=mapa)
    todas = COLUNAS_DI | COLUNAS_PCT | {"liq.2meses", "patrim.liq"}
    for col in todas:
        if col in df.columns:
            df[col] = df[col].apply(
                lambda v: converter_numero_br(v, col in COLUNAS_PCT, col in COLUNAS_DI)
            )
    df = df[df.index.notna() & (df.index.astype(str).str.strip() != "")]
    print(f"  {len(df)} ações obtidas")
    return df


def buscar_ibovespa():
    """Busca composição atual do Ibovespa na B3."""
    url = (
        "https://sistemaswebb3-listados.b3.com.br/indexProxy/indexCall/"
        "GetPortfolioDay/eyJsYW5ndWFnZSI6InB0LWJyIiwicGFnZU51bWJlciI6MSwi"
        "cGFnZVNpemUiOjEyMCwiaW5kZXgiOiJJQk9WIiwic2VnbWVudCI6IjIifQ=="
    )
    try:
        r = requests.get(url, headers={"User-Agent": "Mozilla/5.0"}, timeout=15)
        r.raise_for_status()
        tickers = [i["cod"] for i in r.json()["results"]]
        print(f"  {len(tickers)} ações no Ibovespa")
        return tickers
    except Exception as e:
        print(f"  Fallback lista BESST ({e})")
        return TODOS_TICKERS_BESST


def get_setor(ticker):
    """Retorna o setor B.E.S.S.T. de um ticker."""
    return next((s for s, l in TICKERS_BESST.items() if ticker in l), "Outro")


def aplicar_filtros_ranking(df, tickers_ibov):
    """Aplica filtros do ranking principal (TOP 10)."""
    df = df[df.index.isin(tickers_ibov)].copy()
    for col in ["liq.2meses", "p/l", "dy", "p/vp"]:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors="coerce")
    if "liq.2meses" in df.columns:
        df = df[df["liq.2meses"] >= FILTRO_LIQUIDEZ_MINIMA]
    if "p/l" in df.columns:
        df = df[(df["p/l"] > FILTRO_PL_MINIMO) & (df["p/l"] <= FILTRO_PL_MAXIMO)]
    if "dy" in df.columns:
        df = df[df["dy"] >= FILTRO_DY_MINIMO]
    if "p/vp" in df.columns:
        df = df[df["p/vp"] <= FILTRO_PVPA_MAXIMO]
    if FILTRO_BESST_ATIVO:
        df = df[df.index.isin(TODOS_TICKERS_BESST)]
    df["setor"] = df.index.map(get_setor)
    print(f"  {len(df)} ações após filtros do ranking")
    return df


def aplicar_filtros_oportunidades(df, tickers_ranking_top10):
    """
    Aplica filtros das Oportunidades Ocultas.

    Filosofia: mostrar empresas BESST que NÃO aparecem no ranking
    principal (Seção 1), mas que têm fundamentos suficientes para
    melhorar. São empresas com DY momentaneamente abaixo de 6%,
    ou com P/L fora da faixa do ranking, mas com ROE positivo e
    valuation atrativo — candidatas a acumulação de longo prazo.

    Diferenças em relação ao ranking principal:
    - Busca em TODO o universo BESST, não apenas no Ibovespa
    - Exclui as empresas que já aparecem no TOP 10 (Seção 1)
    - Sem filtro de DY mínimo — captura pagadoras em momento fraco
    - Sem filtro de liquidez — estratégia de acumulação e manutenção
    - P/VP <= 1,5x — descontada ou próxima do valor patrimonial
    - ROE > 8% — ainda operacionalmente sadia

    Parâmetros:
        df                    : DataFrame completo do Fundamentus
        tickers_ranking_top10 : lista de tickers já no TOP 10 — serão excluídos
    """
    # Busca em TODO o universo BESST — sem restrição de Ibovespa
    df = df[df.index.isin(TODOS_TICKERS_BESST)].copy()

    # Remove as empresas que já aparecem na Seção 1
    # O objetivo é mostrar empresas NOVAS, não repetir o ranking
    df = df[~df.index.isin(tickers_ranking_top10)]

    for col in ["p/l", "p/vp", "roe", "cotacao"]:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors="coerce")

    if "roe" in df.columns:
        df = df[df["roe"] >= OO_ROE_MINIMO]
    if "p/vp" in df.columns:
        df = df[df["p/vp"] <= OO_PVPA_MAXIMO]
    if "p/l" in df.columns:
        df = df[df["p/l"] > OO_PL_MINIMO]

    df["setor"] = df.index.map(get_setor)
    print(f"  {len(df)} ações após filtros de oportunidades")
    return df


def calcular_ranking(df):
    """Calcula ranking multicritério com normalização min-max."""
    df = df.copy()
    indicadores = {
        "dy":         ("dy",            True),
        "roe":        ("roe",           True),
        "pl":         ("p/l",           False),
        "pvpa":       ("p/vp",          False),
        "div_ebitda": ("div.liq.patrim",False),
    }
    for met, (col, maior) in indicadores.items():
        if col not in df.columns:
            continue
        s = pd.to_numeric(df[col], errors="coerce").dropna()
        if s.empty or s.max() == s.min():
            df[f"nota_{met}"] = 0.5
            continue
        vmin, vmax = s.min(), s.max()
        df[f"nota_{met}"] = ((df[col]-vmin)/(vmax-vmin)) if maior \
                            else ((vmax-df[col])/(vmax-vmin))
    df["nota_final"] = sum(
        df[f"nota_{m}"].fillna(0) * p
        for m, p in PESOS.items()
        if f"nota_{m}" in df.columns
    )
    df = df.sort_values("nota_final", ascending=False)
    df.insert(0, "ranking", range(1, len(df)+1))
    return df


def buscar_historico_dividendos_anual(tickers):
    """
    Busca dividendos anuais dos últimos N anos via yfinance.
    Retorna DataFrame com dividendo total por ano por ticker,
    e também CAGR e consistência para cada ação.
    """
    data_ini = datetime.today() - timedelta(days=ANOS_HISTORICO * 365)
    anos_ref  = list(range(datetime.today().year - ANOS_HISTORICO,
                           datetime.today().year + 1))

    historico  = {}   # {ticker: {ano: dividendo_total}}
    indicadores = []  # lista de dicts com CAGR e consistência

    for ticker in tickers:
        try:
            divs = yf.Ticker(f"{ticker}.SA").dividends
            if divs.empty:
                historico[ticker] = {a: 0 for a in anos_ref}
                indicadores.append({"ticker": ticker, "cagr_div": float("nan"),
                                    "consistencia_div": 0.0, "anos_com_div": 0})
                time.sleep(0.3)
                continue

            divs.index = divs.index.tz_localize(None)
            divs = divs[divs.index >= pd.Timestamp(data_ini)]

            # Soma dividendos por ano
            da = divs.groupby(divs.index.year).sum()

            # Preenche anos sem pagamento com 0
            for ano in anos_ref:
                if ano not in da.index:
                    da[ano] = 0
            da = da.sort_index()

            historico[ticker] = da.to_dict()

            # Calcula CAGR e consistência
            anos_com_div = int((da > 0).sum())
            consistencia = min(anos_com_div / ANOS_HISTORICO, 1.0)
            da_pos = da[da > 0]
            if len(da_pos) >= 2:
                cagr = (da_pos.iloc[-1] / da_pos.iloc[0]) ** \
                       (1 / (len(da_pos) - 1)) - 1
            else:
                cagr = float("nan")

            indicadores.append({
                "ticker":          ticker,
                "cagr_div":        cagr,
                "consistencia_div":consistencia,
                "anos_com_div":    anos_com_div,
            })
            time.sleep(0.5)

        except Exception as e:
            print(f"  Erro em {ticker}: {e}")
            historico[ticker] = {a: 0 for a in anos_ref}
            indicadores.append({"ticker": ticker, "cagr_div": float("nan"),
                                 "consistencia_div": 0.0, "anos_com_div": 0})
            time.sleep(0.5)

    df_hist = pd.DataFrame(indicadores).set_index("ticker")
    return historico, df_hist, anos_ref


# --- SEÇÃO 4: Funções de gráficos ---

def grafico_pl_pvp(df_top10):
    """
    Gera gráfico de barras duplas (P/L e P/VP) para as TOP 10 ações.
    Retorna a imagem como bytes para inserir no Word.
    """
    tickers = df_top10.index.tolist()
    pl_vals  = [df_top10.loc[t, "p/l"]  if pd.notna(df_top10.loc[t, "p/l"])  else 0 for t in tickers]
    pvp_vals = [df_top10.loc[t, "p/vp"] if pd.notna(df_top10.loc[t, "p/vp"]) else 0 for t in tickers]

    x     = range(len(tickers))
    larg  = 0.35  # Largura de cada barra

    fig, ax = plt.subplots(figsize=(12, 5))
    fig.patch.set_facecolor("#FAFAFA")
    ax.set_facecolor("#FAFAFA")

    barras_pl  = ax.bar([i - larg/2 for i in x], pl_vals,  larg,
                        label="P/L",  color="#1F497D", alpha=0.85)
    barras_pvp = ax.bar([i + larg/2 for i in x], pvp_vals, larg,
                        label="P/VP", color="#70AD47", alpha=0.85)

    # Linha de referência P/VP = 1 (desconto sobre patrimônio)
    ax.axhline(y=1.0, color="#ED7D31", linestyle="--", linewidth=1.2,
               alpha=0.8, label="P/VP = 1,0 (referência)")

    # Rótulos de valor no topo de cada barra
    for barra in barras_pl:
        h = barra.get_height()
        if h > 0:
            ax.text(barra.get_x() + barra.get_width()/2, h + 0.1,
                    f"{h:.1f}", ha="center", va="bottom",
                    fontsize=7.5, color="#1F497D", fontweight="bold")

    for barra in barras_pvp:
        h = barra.get_height()
        if h > 0:
            ax.text(barra.get_x() + barra.get_width()/2, h + 0.1,
                    f"{h:.2f}", ha="center", va="bottom",
                    fontsize=7.5, color="#375623", fontweight="bold")

    ax.set_xticks(list(x))
    ax.set_xticklabels(tickers, rotation=30, ha="right", fontsize=9)
    ax.set_ylabel("Múltiplo (vezes)", fontsize=10)
    ax.set_title("Comparativo P/L e P/VP — TOP 10", fontsize=12,
                 fontweight="bold", color="#1F497D", pad=12)
    ax.legend(fontsize=9, loc="upper right")
    ax.yaxis.set_minor_locator(mticker.AutoMinorLocator())
    ax.grid(axis="y", alpha=0.3, linestyle="--")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)

    plt.tight_layout()

    # Salva em memória como PNG para inserir no Word
    buf = BytesIO()
    plt.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def grafico_dy_historico(historico, tickers_top5, anos_ref):
    """
    Gera gráfico de linhas com evolução do DY histórico das TOP 5.

    O DY é calculado como dividendo anual pago dividido pelo preço
    atual da ação — métrica consistente para comparação entre empresas.
    Retorna a imagem como bytes para inserir no Word.
    """
    fig, ax = plt.subplots(figsize=(12, 5))
    fig.patch.set_facecolor("#FAFAFA")
    ax.set_facecolor("#FAFAFA")

    for i, ticker in enumerate(tickers_top5):
        if ticker not in historico:
            continue
        dados_ano = historico[ticker]
        anos  = sorted(dados_ano.keys())
        divs  = [dados_ano.get(a, 0) for a in anos]

        # Filtra apenas anos dentro do período de análise
        anos_plot = [a for a in anos if a in anos_ref]
        divs_plot = [dados_ano.get(a, 0) for a in anos_plot]

        cor = CORES_GRAFICOS[i % len(CORES_GRAFICOS)]

        ax.plot(anos_plot, divs_plot,
                marker="o", linewidth=2, markersize=6,
                label=ticker, color=cor)

        # Rótulo no último ponto
        if anos_plot and divs_plot:
            ax.annotate(
                f"R$ {divs_plot[-1]:.2f}",
                xy=(anos_plot[-1], divs_plot[-1]),
                xytext=(5, 5), textcoords="offset points",
                fontsize=8, color=cor, fontweight="bold"
            )

    ax.set_xlabel("Ano", fontsize=10)
    ax.set_ylabel("Dividendo Total Pago (R$)", fontsize=10)
    ax.set_title("Evolução dos Dividendos Anuais — TOP 5 (últimos 5 anos)",
                 fontsize=12, fontweight="bold", color="#1F497D", pad=12)
    ax.legend(fontsize=9, loc="upper left")
    ax.set_xticks(anos_ref)
    ax.grid(alpha=0.3, linestyle="--")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)

    plt.tight_layout()

    buf = BytesIO()
    plt.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


# --- SEÇÃO 5: Formatação Word ---

def cel_cabecalho(celula, texto):
    """Formata célula de cabeçalho com fundo azul e texto branco."""
    tc = celula._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "1F497D")
    tcPr.append(shd)
    p = celula.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(texto)
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def cel_dados(celula, texto, negrito=False, centro=True, cor=None):
    """Formata célula de dados."""
    p = celula.paragraphs[0]
    if centro:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(str(texto))
    r.font.size = Pt(9)
    r.bold = negrito
    if cor:
        r.font.color.rgb = cor


def zebra(linha, numero):
    """Aplica fundo cinza alternado nas linhas pares."""
    if numero % 2 == 0:
        for c in linha.cells:
            tc = c._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  "F2F2F2")
            tcPr.append(shd)


def add_par(doc, texto, tam=10, negrito=False, cor=None,
            ea=0, ed=4, recuo=0, italico=False, justificado=False):
    """Adiciona parágrafo formatado ao documento."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(ea)
    p.paragraph_format.space_after  = Pt(ed)
    if recuo:
        p.paragraph_format.left_indent = Cm(recuo)
    if justificado:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(texto)
    r.font.size = Pt(tam)
    r.bold   = negrito
    r.italic = italico
    if cor:
        r.font.color.rgb = cor
    return p


# --- SEÇÃO 6: Geração do documento ---

def gerar_sumario(doc, agora):
    """Gera sumário dinâmico com estrutura do relatório mensal."""
    h = doc.add_heading("SUMARIO", level=1)
    h.runs[0].font.color.rgb = COR_TITULO

    itens = [
        ("Introducao e Objetivos",                        3),
        ("Secao 1 — Ranking TOP 10 e Comparativo Visual", 4),
        ("Secao 2 — Evolucao de Dividendos TOP 5",        5),
        ("Secao 3 — Oportunidades Ocultas",               6),
        ("Siglas e Termos",                               7),
        ("Metadados",                                     8),
    ]
    for titulo, pag in itens:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15))
        r = p.add_run(titulo)
        r.font.size = Pt(10)
        r2 = p.add_run(f"\t{pag}")
        r2.font.size = Pt(10)

    nota = doc.add_paragraph()
    rn = nota.add_run(
        "Nota: numeros de pagina sao estimativas. "
        "Clique com o botao direito no sumario e selecione "
        "Atualizar campo para paginacao exata."
    )
    rn.font.size = Pt(8)
    rn.italic = True
    rn.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def gerar_relatorio_mensal(df_ranking, df_oportunidades,
                            historico_divs, df_hist_ind,
                            anos_ref):
    """
    Função principal que monta o documento Word do relatório mensal.

    Parâmetros:
        df_ranking       : DataFrame com o ranking TOP 10
        df_oportunidades : DataFrame com as oportunidades ocultas filtradas
        historico_divs   : dict {ticker: {ano: dividendo}} das TOP 5
        df_hist_ind      : DataFrame com CAGR e consistência das TOP 5
        anos_ref         : lista de anos do histórico
    """
    doc   = Document()
    agora = datetime.now()

    # Configuração da página A4
    sec = doc.sections[0]
    sec.page_width    = Cm(21.0)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)

    df_top10 = df_ranking.head(TOP_RANKING).copy()
    df_top5  = df_ranking.head(TOP_HISTORICO).copy()

    # ── CAPA ─────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    r = p.add_run("RELATORIO MENSAL DE INVESTIMENTOS")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = COR_TITULO

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Carteira Previdenciaria — Estrategia de Dividendos de Longo Prazo")
    r2.font.size = Pt(14)
    r2.font.color.rgb = COR_SUBTITULO

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(agora.strftime("Referencia: %B de %Y").title())
    r3.font.size = Pt(11)

    doc.add_page_break()

    # ── SUMÁRIO ───────────────────────────────────────────────────────────
    gerar_sumario(doc, agora)
    doc.add_page_break()

    # ── INTRODUÇÃO ────────────────────────────────────────────────────────
    h = doc.add_heading("INTRODUCAO E OBJETIVOS DO RELATORIO", level=1)
    h.runs[0].font.color.rgb = COR_TITULO

    add_par(doc, TEXTO_INTRODUCAO, ed=8, justificado=True)

    for titulo_sec, texto_sec in [
        ("Objetivo",            TEXTO_OBJETIVO),
        ("Metodologia",         TEXTO_METODOLOGIA),
        ("Limitacoes",          TEXTO_LIMITACOES),
        ("Aviso Importante — Nao Constitui Recomendacao de Investimento", TEXTO_AVISO),
    ]:
        add_par(doc, titulo_sec, tam=11, negrito=True,
                cor=COR_SUBTITULO, ea=8, ed=4)
        for bloco in texto_sec.split("\n\n"):
            if bloco.strip():
                add_par(doc, bloco.strip(), ed=6, justificado=True)

    doc.add_page_break()

    # ── SEÇÃO 1: RANKING TOP 10 + GRÁFICO ────────────────────────────────
    h1 = doc.add_heading(
        "SECAO 1 — Ranking TOP 10 e Comparativo Visual", level=1
    )
    h1.runs[0].font.color.rgb = COR_TITULO

    # Parâmetros aplicados
    p_param = doc.add_paragraph()
    rl = p_param.add_run("Filtros: ")
    rl.bold = True
    rl.font.size = Pt(9)
    rv = p_param.add_run(
        f"DY >= {FILTRO_DY_MINIMO:.0%}  |  "
        f"P/L entre {FILTRO_PL_MINIMO} e {FILTRO_PL_MAXIMO}  |  "
        f"P/VP <= {FILTRO_PVPA_MAXIMO}  |  "
        f"Liquidez >= R$ {FILTRO_LIQUIDEZ_MINIMA:,.0f}  |  "
        f"Setores B.E.S.S.T."
    )
    rv.font.size = Pt(9)
    rv.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph()

    # Tabela ranking
    cols = ["#", "Ticker", "Setor", "DY", "P/L", "P/VP", "ROE",
            "CAGR Div", "Consist.", "Nota"]
    tab = doc.add_table(rows=1, cols=len(cols))
    tab.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, c in enumerate(cols):
        cel_cabecalho(tab.rows[0].cells[i], c)

    for _, row in df_top10.iterrows():
        ticker  = row.name
        rank_n  = int(row["ranking"])
        dy_v    = row.get("dy")
        cagr_v  = df_hist_ind.loc[ticker, "cagr_div"] \
                  if ticker in df_hist_ind.index else float("nan")
        cons_v  = df_hist_ind.loc[ticker, "consistencia_div"] \
                  if ticker in df_hist_ind.index else 0.0
        cagr_s  = f"{float(cagr_v):.1%}" \
                  if not math.isnan(float(cagr_v)) else "N/D"
        cons_s  = f"{float(cons_v):.0%}" if cons_v else "N/D"

        nova = tab.add_row()
        vals = [
            rank_n,
            ticker,
            row.get("setor", ""),
            f"{dy_v:.1%}"          if pd.notna(dy_v)              else "N/D",
            f"{row['p/l']:.1f}"    if pd.notna(row.get("p/l"))    else "N/D",
            f"{row['p/vp']:.2f}"   if pd.notna(row.get("p/vp"))   else "N/D",
            f"{row['roe']:.1%}"    if pd.notna(row.get("roe"))     else "N/D",
            cagr_s,
            cons_s,
            f"{row['nota_final']:.3f}",
        ]
        for i, v in enumerate(vals):
            cor = None
            if i == 3 and pd.notna(dy_v):
                cor = COR_DESTAQUE if dy_v >= 0.08 \
                      else RGBColor(0xBF, 0x8F, 0x00)
            cel_dados(nova.cells[i], v,
                      negrito=(i == 1), centro=(i != 2), cor=cor)
        zebra(nova, rank_n)

    add_par(doc,
            "Metodologia: ranking multicriterio min-max. "
            "Pesos: DY 30% | ROE 25% | P/L 20% | P/VP 15% | Div/Patrim 10%.",
            tam=8, italico=True,
            cor=RGBColor(0x80, 0x80, 0x80), ea=6)

    # Gráfico P/L e P/VP
    doc.add_paragraph()
    add_par(doc, "Comparativo de Valuation — P/L e P/VP",
            tam=11, negrito=True, cor=COR_SUBTITULO, ea=6, ed=4)
    add_par(doc,
            "O grafico abaixo compara os multiplos de valuation das 10 acoes "
            "classificadas no ranking. A linha laranja em P/VP = 1,0 indica "
            "o ponto em que a acao negocia ao valor exato do patrimonio liquido. "
            "Acoes abaixo dessa linha negociam com desconto sobre o patrimonio.",
            tam=9, ed=6, justificado=True)

    print("  Gerando gráfico P/L e P/VP...")
    buf_pl_pvp = grafico_pl_pvp(df_top10)
    doc.add_picture(buf_pl_pvp, width=Cm(16))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ── SEÇÃO 2: EVOLUÇÃO DE DIVIDENDOS TOP 5 ────────────────────────────
    h2 = doc.add_heading("SECAO 2 — Evolucao de Dividendos TOP 5", level=1)
    h2.runs[0].font.color.rgb = COR_TITULO

    add_par(doc,
            "O grafico a seguir apresenta o total de dividendos pagos por ano "
            "pelas 5 melhores acoes do ranking nos ultimos 5 anos. "
            "A tendencia de crescimento das linhas indica empresas que aumentam "
            "consistentemente seus proventos — caracteristica essencial para "
            "uma carteira previdenciaria de longo prazo.",
            tam=10, ed=8, justificado=True)

    print("  Gerando gráfico de DY histórico...")
    tickers_top5 = df_top5.index.tolist()
    buf_dy = grafico_dy_historico(historico_divs, tickers_top5, anos_ref)
    doc.add_picture(buf_dy, width=Cm(16))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Tabela resumo histórico TOP 5
    doc.add_paragraph()
    add_par(doc, "Resumo — Historico de Dividendos TOP 5",
            tam=11, negrito=True, cor=COR_SUBTITULO, ea=8, ed=4)

    cols_hist = ["Ticker", "Setor", "DY Atual", "CAGR Div 5a",
                 "Consist.", "Anos c/ Div"]
    tab_h = doc.add_table(rows=1, cols=len(cols_hist))
    tab_h.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, c in enumerate(cols_hist):
        cel_cabecalho(tab_h.rows[0].cells[i], c)

    for pos, (ticker, row) in enumerate(df_top5.iterrows(), 1):
        cagr_v = df_hist_ind.loc[ticker, "cagr_div"] \
                 if ticker in df_hist_ind.index else float("nan")
        cons_v = df_hist_ind.loc[ticker, "consistencia_div"] \
                 if ticker in df_hist_ind.index else 0.0
        anos_v = int(df_hist_ind.loc[ticker, "anos_com_div"]) \
                 if ticker in df_hist_ind.index else 0
        cagr_s = f"{float(cagr_v):.1%}" \
                 if not math.isnan(float(cagr_v)) else "N/D"
        cons_s = f"{float(cons_v):.0%}" if cons_v else "N/D"

        nova = tab_h.add_row()
        vals = [
            ticker,
            row.get("setor", ""),
            f"{row['dy']:.1%}" if pd.notna(row.get("dy")) else "N/D",
            cagr_s,
            cons_s,
            f"{anos_v}/{ANOS_HISTORICO}",
        ]
        # Cor do CAGR: verde se positivo, vermelho se negativo
        for i, v in enumerate(vals):
            cor = None
            if i == 3:  # CAGR
                try:
                    val_f = float(v.replace("%",""))
                    cor = COR_DESTAQUE if val_f > 0 else COR_RISCO
                except Exception:
                    pass
            cel_dados(nova.cells[i], v,
                      negrito=(i == 0), centro=(i != 1), cor=cor)
        zebra(nova, pos)

    doc.add_page_break()

    # ── SEÇÃO 3: OPORTUNIDADES OCULTAS ───────────────────────────────────
    h3 = doc.add_heading("SECAO 3 — Oportunidades Ocultas", level=1)
    h3.runs[0].font.color.rgb = COR_TITULO

    add_par(doc,
            "Esta secao identifica acoes dos setores de renda passiva que "
            "negociam abaixo do seu valor patrimonial (P/VP <= 1,0x) e "
            "apresentam fundamentos operacionais sadios. A logica e simples: "
            "o preco nominal baixo do papel — combinado com desconto sobre o "
            "patrimonio — permite acumular volume expressivo com desembolso "
            "reduzido, posicionando o investidor para capturar a valorizacao "
            "no momento em que o mercado reconhecer o valor real da empresa.",
            tam=10, ed=6, justificado=True)

    add_par(doc, "Criterios de selecao aplicados:",
            tam=10, negrito=True, ea=4, ed=2)

    criterios = [
        "Universo: todos os tickers BESST — sem restricao de indice",
        "Excluidas as empresas ja presentes no ranking TOP 10 (Secao 1)",
        f"P/VP <= {OO_PVPA_MAXIMO:.1f}x — descontada ou proxima do patrimonio",
        f"ROE > {OO_ROE_MINIMO:.0%} — empresa ainda operacionalmente sadia",
        "P/L positivo — sem prejuizo nos ultimos 12 meses",
        "Ordenacao por P/VP crescente (mais descontado primeiro)",
        "Sem filtro de DY minimo — captura pagadoras em momento fraco",
        "Sem filtro de liquidez — estrategia de acumulacao e manutencao",
    ]
    for crit in criterios:
        add_par(doc, f"- {crit}", tam=9, recuo=0.5, ed=2)

    doc.add_paragraph()

    # Busca CAGR para as oportunidades ocultas
    tickers_oo = df_oportunidades.head(OO_TOP).index.tolist()

    # Colunas: adicionamos Cotacao para referência de preço nominal
    cols_oo = ["#", "Ticker", "Setor", "Cotacao", "P/VP",
               "ROE", "DY Atual", "CAGR Div", "P/L"]
    tab_oo = doc.add_table(rows=1, cols=len(cols_oo))
    tab_oo.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, col_nome in enumerate(cols_oo):
        cel_cabecalho(tab_oo.rows[0].cells[i], col_nome)

    for pos, (ticker, row) in enumerate(
            df_oportunidades.head(OO_TOP).iterrows(), 1):
        cagr_v = df_hist_ind.loc[ticker, "cagr_div"] \
                 if ticker in df_hist_ind.index else float("nan")
        cagr_s = f"{float(cagr_v):.1%}" \
                 if not math.isnan(float(cagr_v)) else "N/D"
        cotacao_v = row.get("cotacao")

        nova = tab_oo.add_row()
        vals = [
            pos,
            ticker,
            row.get("setor", ""),
            # Cotação em reais — referência de preço nominal do papel
            f"R$ {cotacao_v:.2f}" if pd.notna(cotacao_v) else "N/D",
            f"{row['p/vp']:.2f}x" if pd.notna(row.get("p/vp")) else "N/D",
            f"{row['roe']:.1%}"   if pd.notna(row.get("roe"))  else "N/D",
            f"{row['dy']:.1%}"    if pd.notna(row.get("dy"))   else "N/D",
            cagr_s,
            f"{row['p/l']:.1f}"   if pd.notna(row.get("p/l"))  else "N/D",
        ]
        for i, v in enumerate(vals):
            cor = None
            # P/VP: verde se <= 0.5 (grande desconto), amarelo se <= 1.0
            if i == 4:
                try:
                    val_f = float(v.replace("x",""))
                    if val_f <= 0.5:
                        cor = COR_DESTAQUE
                    elif val_f <= 1.0:
                        cor = RGBColor(0xBF, 0x8F, 0x00)
                except Exception:
                    pass
            # CAGR: verde se positivo, vermelho se negativo
            if i == 7:
                try:
                    val_f = float(v.replace("%",""))
                    cor = COR_DESTAQUE if val_f > 0 else COR_RISCO
                except Exception:
                    pass
            cel_dados(nova.cells[i], v,
                      negrito=(i == 1), centro=(i != 2), cor=cor)
        zebra(nova, pos)

    add_par(doc,
            "Nota: empresas ordenadas por P/VP crescente. Nenhuma delas aparece "
            "no ranking principal (Secao 1) — sao oportunidades em momento de "
            "desvalorizacao com fundamentos para recuperacao. A cotacao e "
            "apresentada como referencia de aporte. P/VP verde = desconto "
            "expressivo (abaixo de 0,50x); amarelo = desconto moderado (ate 1,0x).",
            tam=8, italico=True,
            cor=RGBColor(0x80, 0x80, 0x80), ea=6)

    doc.add_page_break()

    # ── SIGLAS E TERMOS ───────────────────────────────────────────────────
    h_sig = doc.add_heading(
        "SIGLAS E TERMOS UTILIZADOS NESTE RELATORIO", level=1
    )
    h_sig.runs[0].font.color.rgb = COR_TITULO

    for termo, definicao in SIGLAS_TERMOS:
        add_par(doc, termo, tam=10, negrito=True, ea=6, ed=1)
        add_par(doc, definicao, tam=10, recuo=0.5, ed=2)

    # ── METADADOS ─────────────────────────────────────────────────────────
    doc.add_page_break()
    h_meta = doc.add_heading("Metadados do Relatorio", level=2)
    h_meta.runs[0].font.color.rgb = COR_TITULO

    for k, v in [
        ("Data de geracao",             agora.strftime("%d/%m/%Y %H:%M")),
        ("Referencia",                  agora.strftime("%B/%Y").title()),
        ("Fontes de dados",             "Fundamentus.com.br, B3, yfinance"),
        ("Acoes Ibovespa analisadas",   "~79"),
        ("Acoes no ranking final",      str(len(df_ranking))),
        ("Oportunidades identificadas", str(min(len(df_oportunidades), OO_TOP))),
        ("Filtro DY minimo (ranking)",  f"{FILTRO_DY_MINIMO:.0%}"),
        ("Filtro P/VP (oportunidades)", f"<= {OO_PVPA_MAXIMO:.1f}x — abaixo do patrimonio"),
        ("Filtro DY (oportunidades)", "Sem filtro de DY — foco em valor"),
        ("Filtro liquidez (oportunidades)", "Sem filtro — estrategia buy and hold"),
        ("Filtro BESST",                "Ativo"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        rk = p.add_run(f"{k}: ")
        rk.bold = True
        rk.font.size = Pt(9)
        rv2 = p.add_run(v)
        rv2.font.size = Pt(9)

    add_par(doc,
            "Este relatorio foi gerado automaticamente para fins educacionais. "
            "Nenhuma informacao constitui recomendacao formal de investimento.",
            tam=8, italico=True,
            cor=RGBColor(0x80, 0x80, 0x80), ea=8)

    # ── Salvar ────────────────────────────────────────────────────────────
    nome    = f"relatorio_mensal_besst_{agora.strftime('%Y%m')}.docx"
    caminho = os.path.join(PASTA_RELATORIOS, nome)
    doc.save(caminho)
    return caminho


# --- SEÇÃO 7: Execução principal ---

if __name__ == "__main__":
    inicio = datetime.now()

    print("\n" + "=" * 60)
    print("  RELATORIO MENSAL — CARTEIRA DE DIVIDENDOS")
    print("  Projeto B3 — Analise de Investimentos")
    print("=" * 60)

    # Passo 1: Criar pastas
    criar_pastas()

    # Passo 2: Coletar dados
    print("\nColetando dados de mercado...")
    tickers_ibov = buscar_ibovespa()
    time.sleep(1)
    df_todos = buscar_fundamentus()

    # Passo 3: Ranking principal (TOP 10)
    print("\nAplicando filtros do ranking...")
    df_ranking_filtrado = aplicar_filtros_ranking(df_todos, tickers_ibov)
    if df_ranking_filtrado.empty:
        print("Nenhuma ação passou pelos filtros do ranking. Encerrando.")
        exit(1)
    df_ranking = calcular_ranking(df_ranking_filtrado)

    # Passo 4: Oportunidades Ocultas
    # Passa os tickers do TOP 10 para excluí-los da Seção 3
    tickers_top10 = df_ranking.head(TOP_RANKING).index.tolist()
    print("\nAplicando filtros de oportunidades ocultas...")
    df_oo_filtrado = aplicar_filtros_oportunidades(df_todos, tickers_top10)

    # Passo 5: Histórico de dividendos (TOP 5 do ranking + oportunidades)
    # Busca para todas de uma vez para economizar chamadas ao yfinance
    tickers_hist = list(set(
        df_ranking.head(TOP_HISTORICO).index.tolist() +
        df_oo_filtrado.head(OO_TOP).index.tolist()
    ))
    print(f"\nBuscando histórico de dividendos para {len(tickers_hist)} ações...")
    historico_divs, df_hist_ind, anos_ref = buscar_historico_dividendos_anual(
        tickers_hist
    )

    # Ordena oportunidades ocultas por P/VP crescente
    # As mais descontadas em relação ao patrimônio aparecem primeiro
    df_oo_filtrado = df_oo_filtrado.copy()
    df_oo_filtrado["cagr_div"] = df_oo_filtrado.index.map(
        lambda t: df_hist_ind.loc[t, "cagr_div"]
                  if t in df_hist_ind.index else float("nan")
    )
    df_oo_filtrado = df_oo_filtrado.sort_values(
        "p/vp", ascending=True, na_position="last"
    )

    # Passo 6: Gerar relatório Word
    print("\nGerando relatório Word com gráficos...")
    caminho = gerar_relatorio_mensal(
        df_ranking, df_oo_filtrado,
        historico_divs, df_hist_ind, anos_ref
    )

    fim   = datetime.now()
    tempo = (fim - inicio).seconds

    print("\n" + "=" * 60)
    print(f"  Relatorio gerado em {tempo} segundos!")
    print(f"  {os.path.abspath(caminho)}")
    print("=" * 60 + "\n")
